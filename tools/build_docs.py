#!/usr/bin/env python3
"""
Kétirányú Markdown <-> DOCX konverter a local_artqtml dokumentumaihoz.

A markdown az igazság forrása, de a docx-ből is vissza lehet állítani a markdownt, hogy
egy Wordben szerkesztett példány ne vesszen el. Ezt nem ígéretként tesszük: a `verify`
parancs minden építés után lefuttatja a md -> docx -> md' kört, és eltérés esetén hibával
áll meg.

A KÉSZLET (a "szerződés"). A konverter kizárólag ezeket ismeri:

    # / ## / ###      -> Heading 1 / 2 / 3
    címblokk          -> Title + Subtitle (az első címsor előtti szakasz)
    sima bekezdés     -> Normal
    | a | b | c |     -> táblázat (Table Grid, félkövér fejlécsor)
    **félkövér**      -> inline félkövér
    ![alt](út)        -> beágyazott kép; a kép alt szövege maga a markdown forrás,
                         ezért a visszaút karakterre pontosan helyreállítja
    ```nyelv ... ```  -> kódblokk; soronként egy bekezdés "AIQ Code <nyelv>" stílussal

Minden más -- dőlt, felsorolás-csillaggal, idézet, lábjegyzet, korrektúra, egyesített cella --
KÍVÜL ESIK a készleten. A konverter ilyenkor megáll és megnevezi a helyet; nem dobja el
némán. A csillag szándékosan literális: a dokumentum cron-kifejezéseket tartalmaz
(*/5 * * * *), amiket egy dőlt-parser összetörne.

A KÓDBLOKK a technikai melléklet miatt került be (BL-02): a dokumentum kilenc JSON és PHP
példát tartalmaz, és ezek kerítés nélkül futó szöveggé olvadnak bele a bekezdésekbe -- pontosan
az az importálási sérülés, amit ez a kör kijavít. Két dolog teszi visszafordíthatóvá:

  * a nyelvcímke a STÍLUS NEVÉBEN utazik ("AIQ Code json"), nem a szövegben, ezért nem lehet
    összekeverni a blokk tartalmával;
  * a kódsorokon nincs inline feldolgozás -- a `**` a kódban csillag marad, nem félkövér.

A blokkon belüli üres sor és a behúzás megmarad: a kanonizálás a kerítéseken belülre nem nyúl.

Használat:
    python3 build_docs.py build   <bemenet.md> [-o kimenet.docx]
    python3 build_docs.py reverse <bemenet.docx> [-o kimenet.md]
    python3 build_docs.py verify  <bemenet.md>

Függőség: python-docx  (pip install python-docx)
"""

import argparse
import os
import re
import shutil
import sys
import tempfile

try:
    import docx
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Inches, Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("HIÁNYZÓ FÜGGŐSÉG: pip install python-docx")

HEADING_STYLES = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
LIST_BULLET, LIST_NUMBER = "List Bullet", "List Number"
# A Word "List Paragraph" stílusa nem különbözteti meg a pontozott és a számozott listát -- a
# számozás a bekezdés numPr-jében ül. Beolvasáskor ezért pontozottként vesszük át (a legacy
# dokumentumok ezt használják), kiíráskor viszont a két külön stílust írjuk, amik egyértelműek.
# Az első átvétel emiatt egyszeri, szűk veszteséggel jár: a számozott lista pontozottként jön be,
# és a markdownban kell visszaszámozni. Onnantól a kör pontos.
LIST_PARAGRAPH = "List Paragraph"
STYLE_TO_LEVEL = {v: k for k, v in HEADING_STYLES.items()}
TOC_PLACEHOLDER = "A tartalomjegyzék frissítéséhez: jobb gomb a jegyzéken → Mező frissítése"
IMG_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)$")
BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")
MAX_IMG_W, MAX_IMG_H = Inches(6.0), Inches(7.2)

# A kódblokk nyelvcímkéje a stílus nevében utazik: "AIQ Code" (címke nélkül), "AIQ Code json",
# "AIQ Code php". Így a visszaút a stílusból olvassa ki, és a blokk szövegéhez nem kell nyúlni.
CODE_STYLE = "AIQ Code"
FENCE_OPEN_RE = re.compile(r"^```([A-Za-z0-9_+-]*)$")
FENCE_CLOSE = "```"


def code_style_name(lang):
    return f"{CODE_STYLE} {lang}" if lang else CODE_STYLE


def code_style_lang(style_name):
    """'AIQ Code json' -> 'json'; 'AIQ Code' -> ''. Nem kódstílusra None."""
    if style_name == CODE_STYLE:
        return ""
    if style_name.startswith(CODE_STYLE + " "):
        return style_name[len(CODE_STYLE) + 1:]
    return None


def ensure_code_style(doc, lang):
    """A kódstílus létrehozása, ha még nincs a dokumentumban."""
    name = code_style_name(lang)
    existing = {s.name for s in doc.styles}
    if name in existing:
        return name
    style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = doc.styles["Normal"]
    style.font.name = "Consolas"
    style.font.size = Pt(9)
    # Sorköz nélkül, hogy a blokk egybefüggő szövegtömbként jelenjen meg, ne szellős
    # bekezdéssorozatként -- egy JSON séma így olvasható.
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    return name


def split_fenced(text):
    """[(kódblokk-e, sorok), ...] -- a kerítéssorok a kódblokk részei maradnak."""
    out, cur, in_code = [], [], False
    for line in text.split("\n"):
        if not in_code and FENCE_OPEN_RE.match(line.rstrip()):
            if cur:
                out.append((False, cur))
            cur, in_code = [line.rstrip()], True
            continue
        if in_code:
            cur.append(line.rstrip())
            if line.rstrip() == FENCE_CLOSE:
                out.append((True, cur))
                cur, in_code = [], False
            continue
        cur.append(line)
    if cur:
        # Lezáratlan kerítés: a hibát a md_to_docx nevezi meg, itt prózaként megy tovább.
        out.append((False, cur))
    return out


def canonicalise(text):
    """Kanonikus markdown: blokkonként pontosan egy üres sor, záró szóköz nincs.

    A docx nem tárolja, hogy két bekezdés között volt-e üres sor -- a bekezdés-elválasztás
    nála implicit. Ha a forrás nem kanonikus, a visszaút szükségszerűen eltérne tőle. Ezért
    a forrást hozzuk kanonikus alakra, nem az összehasonlítást lazítjuk.

    A kódblokkokba nem nyúlunk: ott az üres sor és a behúzás a tartalom része, nem tördelés.
    """
    blocks = []
    for is_code, lines in split_fenced(text):
        if is_code:
            blocks.append(lines)
        else:
            blocks.extend(_prose_blocks("\n".join(lines)))
    return "\n\n".join("\n".join(b) for b in blocks) + "\n"


def _prose_blocks(text):
    """A kódblokkokon kívüli szöveg blokkokra bontása (a kanonizálás eredeti logikája)."""
    # Két listaelem (vagy két táblázatsor) közötti üres sor nem hordoz információt, és a docx nem
    # is tudja tárolni: a bekezdés-elválasztás ott implicit. Ha benne hagynánk, a visszaút
    # szükségszerűen eltérne a forrástól. Ezért előbb kiszedjük.
    lines = text.split("\n")
    ROW = re.compile(r"^(\||- |\d+\. )")
    cleaned = []
    for n, raw in enumerate(lines):
        if not raw.strip():
            prev = next((l for l in reversed(cleaned) if l.strip()), "")
            nxt = next((l for l in lines[n + 1:] if l.strip()), "")
            if ROW.match(prev) and ROW.match(nxt):
                continue
        cleaned.append(raw)
    text = "\n".join(cleaned)

    blocks, cur, in_table = [], [], False
    for raw in text.split("\n"):
        line = raw.rstrip()
        is_row = line.startswith("|") or bool(re.match(r"^(- |\d+\. )", line))
        if not line.strip():
            if cur:
                blocks.append(cur); cur = []
            in_table = False
            continue
        if is_row and in_table:
            cur.append(line)
        else:
            if cur:
                blocks.append(cur)
            cur, in_table = [line], is_row
    if cur:
        blocks.append(cur)
    return blocks


class ContractError(Exception):
    """A dokumentum olyat tartalmaz, ami kívül esik a támogatott készleten."""


# --------------------------------------------------------------------------- markdown -> docx

def split_inline(text):
    """'a **b** c' -> [('a ', False), ('b', True), (' c', False)]"""
    out, pos = [], 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            out.append((text[pos:m.start()], False))
        out.append((m.group(1), True))
        pos = m.end()
    if pos < len(text):
        out.append((text[pos:], False))
    return out or [("", False)]


def add_runs(par, text):
    for chunk, bold in split_inline(text):
        if chunk:
            run = par.add_run(chunk)
            run.bold = bold


def add_toc_field(doc):
    par = doc.add_paragraph()
    for kind, payload in (("begin", None), ("instr", ' TOC \\o "1-3" \\h \\z \\u '),
                          ("separate", None), ("text", TOC_PLACEHOLDER), ("end", None)):
        run = par.add_run()
        if kind == "instr":
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve")
            el.text = payload
        elif kind == "text":
            el = OxmlElement("w:t")
            el.text = payload
        else:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), kind)
        run._r.append(el)
    return par


def set_alt_text(shape, value):
    """A kép leírásába (alt szöveg) a markdown forrást írjuk -- ez a visszaút kulcsa."""
    docpr = shape._inline.docPr
    docpr.set("descr", value)
    docpr.set("name", value)


def set_column_widths(table, rows):
    """Karakterbüdzsé-alapú oszlopszélesség.

    Két dolgot kell egyszerre teljesíteni. Egy azonosító (List-003) nem tördelhető, tehát a
    saját oszlopa nem lehet nála keskenyebb -- ez a garantált minimum. Ami a minimumok után
    marad, azt a tényleges tartalomhossz arányában osztjuk szét, így a hosszú Elvárás szöveg
    kapja a helyet, nem az egyenlő elosztás.

    A puszta arányos elosztás mindkét irányban elromlik: egyenlő szélességnél az Elvárás
    oszlop háromsorosra tördelődik, tisztán tartalomarányosnál pedig az azonosító tört szét
    "Lis t- 00 3" alakúra.
    """
    ncols = len(rows[0])
    BUDGET = 95                       # kb. ennyi karakter fér ki 10 pt-tel a szövegtükörben

    need, want = [], []
    for c in range(ncols):
        longest_word = max((len(w) for r in rows for w in r[c].split()), default=1)
        lengths = sorted(len(r[c]) for r in rows)
        p90 = lengths[int(len(lengths) * 0.9)] if lengths else 1
        need.append(min(longest_word + 2, 24))
        want.append(max(p90, longest_word))

    spare = BUDGET - sum(need)
    if spare > 0:
        total_want = sum(want) or 1
        chars = [need[c] + spare * want[c] / total_want for c in range(ncols)]
    else:
        chars = list(need)
    total_chars = sum(chars)

    usable = Inches(6.3)
    table.autofit = False

    # A cellaszélességet a Word és a LibreOffice is figyelmen kívül hagyja, amíg a táblázat
    # elrendezése "auto" - fixre kell állítani, és a táblázat saját szélességét is megadni.
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(int(usable / 635)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)

    for c in range(ncols):
        width = int(usable * chars[c] / total_chars)
        table.columns[c].width = width
        for row in table.rows:
            row.cells[c].width = width


def parse_table_block(lines, i):
    """Visszaadja (sorok, következő index) a | ... | blokkra."""
    rows = []
    while i < len(lines) and lines[i].startswith("|"):
        raw = lines[i].strip()
        cells = [c.strip() for c in raw.strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def md_to_docx(md_path, docx_path):
    lines = open(md_path, encoding="utf-8").read().split("\n")
    doc = docx.Document()
    base = os.path.dirname(os.path.abspath(md_path))

    # --- címblokk: az első címsorig
    first_heading = next((n for n, l in enumerate(lines) if l.startswith("# ")), len(lines))
    front = [l for l in lines[:first_heading] if l.strip()]
    for n, line in enumerate(front):
        text = BOLD_RE.sub(r"\1", line) if n == 0 else line
        par = doc.add_paragraph(style="Title" if n == 0 else "Subtitle")
        add_runs(par, text) if n else par.add_run(text)

    doc.add_paragraph("Tartalomjegyzék", style="Heading 1")
    add_toc_field(doc)

    i = first_heading
    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        if not stripped.strip():
            i += 1
            continue

        m = FENCE_OPEN_RE.match(stripped)
        if m:
            lang = m.group(1)
            style = ensure_code_style(doc, lang)
            start = i
            i += 1
            body = []
            while i < len(lines) and lines[i].rstrip() != FENCE_CLOSE:
                body.append(lines[i].rstrip())
                i += 1
            if i >= len(lines):
                raise ContractError(f"{md_path}:{start+1}: lezáratlan kódblokk")
            i += 1  # a záró kerítés
            for text in body:
                par = doc.add_paragraph(style=style)
                # Szándékosan NEM add_runs: a kódban a ** csillag marad, nem félkövér, és a
                # behúzás a tartalom része. A python-docx a vezető szóközt megőrzi.
                par.add_run(text)
            if not body:
                # Üres blokk is blokk: enélkül a visszaút nem tudná, hogy volt itt kerítés.
                doc.add_paragraph(style=style).add_run("")
            doc.add_paragraph()
            continue

        if stripped.startswith("|"):
            rows, i = parse_table_block(lines, i)
            if not rows:
                continue
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            set_column_widths(table, rows)
            for r, cells in enumerate(rows):
                if len(cells) != len(rows[0]):
                    raise ContractError(f"{md_path}: egyenetlen táblázatsor a {i}. sor körül")
                for c, cell in enumerate(cells):
                    par = table.cell(r, c).paragraphs[0]
                    add_runs(par, cell.replace("\\|", "|"))
                    # A fejlécsor félkövér -- de csak akkor, ha van mit fejlécnek tekinteni. Egy
                    # egysoros táblázat keretes megjegyzés, nem fejléces táblázat: ott a saját
                    # félkövérjét kell megőrizni, különben a "**Fontos**" kiemelés eltűnik.
                    if r == 0 and len(rows) > 1:
                        for run in par.runs:
                            run.bold = True
            doc.add_paragraph()
            continue

        m = re.match(r"^(#{1,6}) (.*)$", stripped)
        if m:
            level = len(m.group(1))
            if level > 3:
                raise ContractError(f"{md_path}:{i+1}: a készlet csak H1-H3 szintet ismer")
            par = doc.add_paragraph(style=HEADING_STYLES[level])
            add_runs(par, m.group(2))
            i += 1
            continue

        m = IMG_RE.match(stripped)
        if m:
            path = os.path.join(base, m.group(2))
            if not os.path.exists(path):
                raise ContractError(f"{md_path}:{i+1}: hiányzó képfájl: {m.group(2)}")
            # Natív mérettel szúrjuk be, majd arányosan zsugorítjuk a szövegtükörbe. Fix
            # szélességgel egy álló képernyőkép (5.5 x 8.5 hüvelyk) 9 hüvelyk magas lenne, ami
            # magasabb a lapnál, és a Word a következő oldalra tolná -- a kép és a hozzá tartozó
            # felirat így elszakadna egymástól.
            doc.add_picture(path)
            shape = doc.inline_shapes[-1]
            scale = min(MAX_IMG_W / shape.width, MAX_IMG_H / shape.height, 1.0)
            shape.width, shape.height = int(shape.width * scale), int(shape.height * scale)
            set_alt_text(shape, stripped)
            i += 1
            continue

        m = re.match(r"^- (.*)$", stripped)
        if m:
            par = doc.add_paragraph(style=LIST_BULLET)
            add_runs(par, m.group(1))
            i += 1
            continue

        m = re.match(r"^\d+\. (.*)$", stripped)
        if m:
            par = doc.add_paragraph(style=LIST_NUMBER)
            add_runs(par, m.group(1))
            i += 1
            continue

        for bad, name in ((r"^\s*[*+] ", "csillagos/plusz felsorolás"),
                          (r"^>", "idézet"), (r"^```", "elrontott kódblokk-kerítés")):
            if re.match(bad, stripped):
                raise ContractError(f"{md_path}:{i+1}: {name} kívül esik a készleten")

        par = doc.add_paragraph(style="Normal")
        add_runs(par, stripped)
        i += 1

    doc.save(docx_path)
    return docx_path


# --------------------------------------------------------------------------- docx -> markdown

LENIENT = {"on": False, "italics": 0, "images": 0}


def runs_to_md(par):
    out = []
    for run in par.runs:
        text = run.text
        if not text:
            continue
        if run.italic:
            # A szerződés nem ismer dőltet: a specifikációban a csillag cron-kifejezéseket jelöl
            # (*/5 * * * *), az aláhúzás pedig azonosítókban szerepel (local_artqtml_log), tehát
            # egyik jelölés sem foglalható le kétértelműség nélkül. Legacy átvételnél ezért a dőlt
            # formázást elhagyjuk -- de megszámoljuk és jelentjük, nem csendben.
            if not LENIENT["on"]:
                raise ContractError(f"dőlt szöveg a készleten kívül: {text[:40]!r}")
            LENIENT["italics"] += 1
        out.append(f"**{text}**" if run.bold else text)
    merged = "".join(out)
    return re.sub(r"\*\*\*\*", "", merged)


def is_toc_paragraph(par):
    return par._p.findall(qn("w:r") + "/" + qn("w:fldChar")) or TOC_PLACEHOLDER in par.text


def docx_to_md(docx_path, md_path):
    doc = docx.Document(docx_path)
    body, out, seen_first_heading = doc.element.body, [], False
    # A Word "List Number" stílusa nem tárolja a látható sorszámot -- azt a szövegszerkesztő
    # számolja. A visszaúton ezért újraszámozzuk a szomszédos elemeket; a számláló minden más
    # blokknál nullázódik. Így az "1. 2. 3." pontosan visszajön, és nem lesz belőle három "1.".
    number = 0
    tables = iter(doc.tables)
    # A kódblokk bekezdésenként érkezik vissza; a kerítést nekünk kell köré tenni, és csak akkor
    # lezárni, amikor a stílus véget ér. A nyelvcímke a stílus nevéből jön.
    code_lang = None

    def close_code():
        nonlocal code_lang
        if code_lang is not None:
            out.append(FENCE_CLOSE)
            out.append("")
            code_lang = None

    for child in body.iterchildren():
        if child.tag == qn("w:tbl"):
            close_code()
            table = next(tables)
            number = 0
            widths = len(table.columns)
            for r, row in enumerate(table.rows):
                cells = []
                for cell in row.cells:
                    text = " ".join(runs_to_md(p) for p in cell.paragraphs).strip()
                    if r == 0 and len(table.rows) > 1:
                        text = re.sub(r"^\*\*(.*)\*\*$", r"\1", text)
                    cells.append(text.replace("|", "\\|"))
                out.append("| " + " | ".join(cells) + " |")
                if r == 0:
                    out.append("|" + "---|" * widths)
            out.append("")
            continue

        if child.tag != qn("w:p"):
            continue
        par = docx.text.paragraph.Paragraph(child, doc)
        style = par.style.name if par.style is not None else "Normal"

        lang = code_style_lang(style)
        if lang is not None:
            if code_lang != lang:
                close_code()
                out.append(FENCE_CLOSE + lang)
                code_lang = lang
            # A kódsor nyers szöveg: se félkövér-jelölés, se csillag-értelmezés.
            out.append(par.text)
            number = 0
            continue
        close_code()

        if is_toc_paragraph(par) or (style == "Heading 1" and par.text.strip() == "Tartalomjegyzék"
                                     and not seen_first_heading):
            continue

        blips = child.findall(".//" + qn("a:blip"))
        if blips:
            descr = child.find(".//" + qn("wp:docPr"))
            src = descr.get("descr") if descr is not None else None
            if not src:
                if not LENIENT["on"]:
                    raise ContractError(
                        "beágyazott kép alt szöveg nélkül -- a forrás nem állítható helyre")
                LENIENT["images"] += 1
                continue
            out.append(src)
            out.append("")
            continue

        text = runs_to_md(par)
        if style != LIST_NUMBER:
            number = 0

        if style in STYLE_TO_LEVEL:
            seen_first_heading = True
            out.append("#" * STYLE_TO_LEVEL[style] + " " + text)
            out.append("")
        elif style in ("Title", "Subtitle"):
            out.append(f"**{text}**" if style == "Title" else text)
            out.append("")
        elif style in (LIST_BULLET, LIST_PARAGRAPH):
            # Nincs záró üres sor: a kanonikus alak szerint a szomszédos listaelemek egyetlen
            # blokkot alkotnak, és a canonicalise() teszi ki az elválasztást a lista után.
            out.append("- " + text)
        elif style == LIST_NUMBER:
            number += 1
            out.append(f"{number}. " + text)
        elif style == "Normal":
            out.append(text)
            out.append("")
        else:
            raise ContractError(f"ismeretlen stílus a készleten kívül: {style!r} -- {par.text[:50]!r}")

    close_code()
    md = canonicalise("\n".join(out))
    open(md_path, "w", encoding="utf-8").write(md)
    return md_path


# --------------------------------------------------------------------------- kör-ellenőrzés

def normalise(text):
    """A körhöz jelentéktelen különbségek kiegyenlítése: záró szóköz, üres sorok száma."""
    lines = [l.rstrip() for l in text.split("\n")]
    return re.sub(r"\n{3,}", "\n\n", "\n".join(lines)).strip()


def verify(md_path, canonicalise_source=True):
    original = open(md_path, encoding="utf-8").read()
    canon = canonicalise(original)
    if canonicalise_source and canon != original:
        open(md_path, "w", encoding="utf-8").write(canon)
        print(f"a forrás kanonikus alakra hozva: {os.path.basename(md_path)}")
    # Az ideiglenes fájlok NEM a forrás mellé kerülnek: az a git-munkakönyvtárba szemetelne,
    # és egy meghiúsult takarítás ott hagyná őket a repóban.
    tmpdir = tempfile.mkdtemp(prefix="build_docs_")
    tmp_docx = os.path.join(tmpdir, "roundtrip.docx")
    tmp_md = os.path.join(tmpdir, "roundtrip.md")
    try:
        md_to_docx(md_path, tmp_docx)
        docx_to_md(tmp_docx, tmp_md)
        a = open(md_path, encoding="utf-8").read().strip()
        b = open(tmp_md, encoding="utf-8").read().strip()
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)

    if a == b:
        print(f"KÖR RENDBEN: {os.path.basename(md_path)} -> docx -> markdown, eltérés nincs")
        return True
    import difflib
    diff = list(difflib.unified_diff(a.split("\n"), b.split("\n"),
                                     "eredeti", "visszaalakított", lineterm="", n=1))
    print(f"KÖR ELTÉRÉS ({len([d for d in diff if d[:1] in '+-' and d[:3] not in ('---', '+++')])} sor):")
    print("\n".join(diff[:60]))
    return False


def main():
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("command", choices=["build", "reverse", "verify"])
    ap.add_argument("input")
    ap.add_argument("-o", "--output")
    ap.add_argument("--legacy", action="store_true",
                    help="egyszeri átvétel régi dokumentumból: a dőlt formázást és az alt szöveg "
                         "nélküli képeket elhagyja, és megszámolva jelenti")
    args = ap.parse_args()
    LENIENT["on"] = args.legacy

    try:
        if args.command == "build":
            out = args.output or os.path.splitext(args.input)[0] + ".docx"
            src = open(args.input, encoding="utf-8").read()
            if canonicalise(src) != src:
                open(args.input, "w", encoding="utf-8").write(canonicalise(src))
                print(f"a forrás kanonikus alakra hozva: {os.path.basename(args.input)}")
            md_to_docx(args.input, out)
            print(f"kész: {out}")
            if not verify(args.input):
                sys.exit("A kör-ellenőrzés elbukott -- a docx nem hű a markdownhoz.")
        elif args.command == "reverse":
            out = args.output or os.path.splitext(args.input)[0] + ".from-docx.md"
            docx_to_md(args.input, out)
            print(f"kész: {out}")
            if LENIENT["on"]:
                print(f"legacy átvétel — elhagyva: {LENIENT['italics']} dőlt szakasz, "
                      f"{LENIENT['images']} kép alt szöveg nélkül")
        else:
            sys.exit(0 if verify(args.input) else 1)
    except ContractError as exc:
        sys.exit(f"KÉSZLETEN KÍVÜLI SZERKEZET: {exc}")


if __name__ == "__main__":
    main()
